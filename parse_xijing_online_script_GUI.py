import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import tkinter as tk
from tkinter import messagebox, filedialog
import os

def parse_script(script_id, save_path):
    url = f'https://aipiaxi.com/article-detail/{script_id}'  # 拼接URL

    try:
        response = requests.get(url)

        # 检查是否错误
        if response.status_code != 200:
            messagebox.showerror("错误", f"网页访问出错，状态码{response.status_code}")
            return

        response.raise_for_status()  # 确保请求成功
        html_content = response.text

        if "timeout:" in html_content:  # 检查页面内容是否为空
            messagebox.showerror("错误", "访问超时，请确认剧本号是否存在")
            return

        # 使用 BeautifulSoup 解析 HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # 创建 Word 文档
        doc = Document()

        is_header = True

        # font = u'等线'  # 设置全局统一字体，目前不可用
        
        # 提取剧本内容
        for paragraph in soup.find_all(['p', 'h2', 'h3', 'h4']):
            
            if paragraph.name != 'p':  # 解析标题类
                # 添加标题部分
                header_text = paragraph.get_text().strip()
                if is_header:  # 第一个出现的标题类必为剧本标题
                    title = doc.add_paragraph(style='Title')
                    run_title = title.add_run(header_text.strip())
                    run_title.font.size = Pt(20)
                    run_title.font.color.rgb = RGBColor(0, 0, 0)
                    run_title.bold = True
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    is_header = False
                else:  # 其余为次级标题或其它
                    heading = doc.add_paragraph(style='Heading' + paragraph.name[-1])
                    run_heading = heading.add_run(header_text.strip())
                    run_heading.font.size = Pt(8 + 2 * eval(paragraph.name[-1]))
                    run_heading.font.color.rgb = RGBColor(0, 0, 0)
                    run_heading.bold = True
                    
            else:  # 解析段落类
                style = {}
                if paragraph.get('style') != None:
                    for i in paragraph.get('style').split(';'):  # 将style字符串构建成字典
                        if i != '':
                            style[i.split(':')[0].strip()] = i.split(':')[1].strip()

                    color = style['color'] if 'color' in style.keys() else '#000000'  # 获取刷色色号
                
                if (paragraph.get('type') == 'conversation') or (paragraph.get('data-number') != None): # 解析对话
                    
                    coversation = doc.add_paragraph(style='Normal')
                    for element in paragraph.children:
                        if ((element.name == 'span' and element.get('type') == 'xj-tips') or element.name == 'strong'):
                            # 加粗部分
                            try:
                                run_strong = coversation.add_run(element.get_text().strip())
                                run_strong.font.color.rgb = RGBColor(0, 0, 0)
                                run_strong.font.size = Pt(11)
                                run_strong.font.bold = True
                            except Exception as e:
                                print("Error 1:", str(e))
                        else:
                            try:
                                # 添加其他文本
                                run_dialog = coversation.add_run(element.string.strip())
                                # 注意RGB两种表示
                                if color[0] == '#':
                                    run_dialog.font.color.rgb = RGBColor(int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))
                                else:
                                    run_dialog.font.color.rgb = RGBColor(eval(color.split(',')[0][-1]), eval(color.split(',')[1]), eval(color.split(',')[2][0]))
                                run_dialog.font.size = Pt(11)
                                run_dialog.font.bold = False
                            except Exception as e:
                                print("Error 2:", str(e))
                            

                else:  # 解析其余
                    special_label = paragraph.find('span', {'type': 'special-label'})
                    if 'text-align' in style.keys():
                        align_style = style['text-align']
                    if special_label:
                        try:
                            p = doc.add_paragraph(special_label.get_text().strip(), style='Normal')
                            if len(p.runs) != 0:
                                run_others = p.runs[0]
                                run_others.font.size = Pt(11)
                                run_others.font.color.rgb = RGBColor(0, 0, 0)
                                run_others.font.bold = True
                        except Exception as e:
                            print("Error 3:", str(e))
                    else:
                        try:
                            p = doc.add_paragraph(paragraph.get_text().strip(), style='Normal')
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            if len(p.runs) != 0:
                                run_center = p.runs[0]
                                run_center.font.size = Pt(11)
                                run_center.font.color.rgb = RGBColor(0, 0, 0)
                                run_center.font.bold = True
                        except Exception as e:
                            print("Error 4:", str(e))
                        
                    

        # 保存 Word 文档
        doc.save(save_path)
        messagebox.showinfo("成功", f"Word文档已生成：{save_path}")

    except Exception as e:
        messagebox.showerror("错误", str(e))

def on_choose_path():
    # 选择文件保存路径
    path = filedialog.askdirectory()  # 选择目录而不是文件
    if path:
        entry_path.delete(0, tk.END)  # 清空当前路径
        entry_path.insert(0, path)  # 插入选择的路径

def on_confirm():
    # 此处添加确认下载逻辑
    script_id = entry_id.get().strip()
    save_path = entry_path.get().strip()
    file_name = entry_file_name.get().strip()

    if not save_path:
        messagebox.showerror("错误", "保存路径未填写！")
        return
    
    # 使用用户输入的文件名或默认文件名
    if file_name:
        full_file_name = f"{file_name}.docx"  # 如果有输入文件名，添加扩展名
    else:
        full_file_name = f"script_{script_id}.docx" if script_id else "script_default.docx"

    full_path = os.path.join(save_path, full_file_name)

    # 这里可以添加你想要执行下载的逻辑
    # 例如：
    messagebox.showinfo("下载确认", f"剧本将保存到：{full_path}")
    parse_script(script_id, full_path)  # 假设这是实际下载函数

# 创建主窗口
root = tk.Tk()
root.title("戏鲸剧本下载器")

# 设置窗口大小
root.geometry("600x250")

# 创建标签和输入框的框架
frame_id = tk.Frame(root)
frame_id.pack(pady=10)

label_id = tk.Label(frame_id, text="戏鲸剧本号：")
label_id.pack(side=tk.LEFT)
label_id.focus_set()

entry_id = tk.Entry(frame_id, width=50)
entry_id.pack(side=tk.LEFT, padx=(5, 0))

# 输入文件名的框架
frame_file_name = tk.Frame(root)
frame_file_name.pack(pady=10)

label_file_name = tk.Label(frame_file_name, text="保存文件名：")
label_file_name.pack(side=tk.LEFT)

entry_file_name = tk.Entry(frame_file_name, width=50)  
entry_file_name.pack(side=tk.LEFT, padx=(5, 0))

# 创建保存路径的框架
frame_path = tk.Frame(root)
frame_path.pack(pady=10)

label_path = tk.Label(frame_path, text="保存路径：")
label_path.pack(side=tk.LEFT)

entry_path = tk.Entry(frame_path, width=40)  # 调整宽度以适应按钮
entry_path.pack(side=tk.LEFT, padx=(5, 0))  # 右侧留出间距

# 选择路径按钮
choose_button = tk.Button(frame_path, text="选择保存路径", command=on_choose_path)
choose_button.pack(side=tk.LEFT)

# 设置初始保存路径为桌面
default_save_path = os.path.join(os.path.expanduser("~"), "Desktop")
entry_path.insert(0, default_save_path)  # 将默认路径插入输入框，不包含文件名

# 确认下载按钮
confirm_button = tk.Button(root, text="确认", command=on_confirm)
confirm_button.pack(pady=20)

root.bind('<Return>', lambda event: on_confirm())

# 启动主循环
root.mainloop()
