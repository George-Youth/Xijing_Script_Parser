import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import messagebox

def fetch_script(script_id):
    url = f'https://aipiaxi.com/article-detail/{script_id}'  # 拼接URL
    try:
        response = requests.get(url)

        # 检查是否错误
        if response.status_code != 200:
            messagebox.showerror("错误", f"网页访问出错，状态码{response.status_code}")
            return

        response.raise_for_status()  # 确保请求成功
        html_content = response.text

        if "timeout:" in html_content:  # 检查页面内容是否为空
            messagebox.showerror("错误", "访问超时，请确认剧本号是否存在")
            return

        return html_content
    
    except Exception as e:
        messagebox.showerror("错误", str(e))

def parse_script(html_content, save_path):

    # 使用 BeautifulSoup 解析 HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # 创建 Word 文档
    doc = Document()

    is_header = True

    # font = u'等线'  # 设置全局统一字体，目前不可用
        
    # 提取剧本内容
    for paragraph in soup.find_all(['p', 'h2', 'h3', 'h4']):
            
        if paragraph.name != 'p':  # 解析标题类
            # 添加标题部分
            header_text = paragraph.get_text().strip()
            if is_header:  # 第一个出现的标题类必为剧本标题
                title = doc.add_paragraph(style='Title')
                run_title = title.add_run(header_text.strip())
                run_title.font.size = Pt(20)
                run_title.font.color.rgb = RGBColor(0, 0, 0)
                run_title.bold = True
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                is_header = False
            else:  # 其余为次级标题或其它
                heading = doc.add_paragraph(style='Heading ' + paragraph.name[-1])
                run_heading = heading.add_run(header_text.strip())
                run_heading.font.size = Pt(8 + 2 * eval(paragraph.name[-1]))
                run_heading.font.color.rgb = RGBColor(0, 0, 0)
                run_heading.bold = True
                    
        else:  # 解析段落类
            style = {}
            if paragraph.get('style') != None:
                for i in paragraph.get('style').split(';'):  # 将style字符串构建成字典
                    if i != '':
                        style[i.split(':')[0].strip()] = i.split(':')[1].strip()

                color = style['color'] if 'color' in style.keys() else '#000000'  # 获取刷色色号
                
            if (paragraph.get('type') == 'conversation') or (paragraph.get('data-number') != None): # 解析对话
                    
                coversation = doc.add_paragraph(style='Normal')
                for element in paragraph.children:
                    if ((element.name == 'span' and element.get('type') == 'xj-tips') or element.name == 'strong'):
                        # 加粗部分
                        try:
                            run_strong = coversation.add_run(element.get_text().strip())
                            run_strong.font.color.rgb = RGBColor(0, 0, 0)
                            run_strong.font.size = Pt(11)
                            run_strong.font.bold = True
                        except Exception as e:
                            print("Error 1:", str(e))
                    else:
                        try:
                            # 添加其他文本
                            run_dialog = coversation.add_run(element.string.strip())
                            # 注意RGB两种表示
                            if color[0] == '#':
                                run_dialog.font.color.rgb = RGBColor(int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))
                            else:
                                run_dialog.font.color.rgb = RGBColor(eval(color.split(',')[0][-1]), eval(color.split(',')[1]), eval(color.split(',')[2][0]))
                            run_dialog.font.size = Pt(11)
                            run_dialog.font.bold = False
                        except Exception as e:
                            print("Error 2:", str(e))
                            

            else:  # 解析其余
                special_label = paragraph.find('span', {'type': 'special-label'})
                if 'text-align' in style.keys():
                    align_style = style['text-align']
                if special_label:
                    try:
                        p = doc.add_paragraph(special_label.get_text().strip(), style='Normal')
                        if len(p.runs) != 0:
                            run_others = p.runs[0]
                            run_others.font.size = Pt(11)
                            run_others.font.color.rgb = RGBColor(0, 0, 0)
                            run_others.font.bold = True
                    except Exception as e:
                        print("Error 3:", str(e))
                else:
                    try:
                        p = doc.add_paragraph(paragraph.get_text().strip(), style='Normal')
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if len(p.runs) != 0:
                            run_center = p.runs[0]
                            run_center.font.size = Pt(11)
                            run_center.font.color.rgb = RGBColor(0, 0, 0)
                            run_center.font.bold = True
                    except Exception as e:
                        print("Error 4:", str(e))


    # 保存 Word 文档
    doc.save(save_path)
    messagebox.showinfo("成功", f"Word文档已生成：{save_path}")
